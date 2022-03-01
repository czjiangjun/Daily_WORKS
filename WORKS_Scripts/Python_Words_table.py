from docx import Document
from docx.shared import Inches

# 打开现有word文档对象
# doc = Document('数据表-3.docx')
doc = Document('2021_04-06-炉号.docx')

# 创建word文档对象
document = Document()
# 添加标题
document.add_heading('标准化数据表-Wang_2021-04-06', 0)


#现有文档的每一段的内容
for para in doc.paragraphs:
#    print(para.text)
# 添加段落
    p = document.add_paragraph(para.text)

#每一段的编号、内容
# for i in range(len(doc.paragraphs)):
#    print(str(i),  doc.paragraphs[i].text)

def Atom_Element(n,i):
#    if j > 0 :
     if (n%3) == 2:
#    Elemets= [58.69, 26.98, 52.00, 58.93, 95.94, 186.21, 180.95, 183.85, 101.07]
        Elemets= [58.69, 26.98, 52.00, 58.93, 95.94, 183.85, 180.95, 186.21, 101.07, 47.88]
     else:
        Elemets= [58.69, 26.98, 52.00, 58.93, 95.94, 183.85, 180.95, 186.21, 101.07, 92.91]
    #Elemets= [58.69, 26.98, 52.00, 58.93, 95.94, 183.85, 180.95, 186.21, 101.07, 178.49]
#    else:
#    Elemets= [58.69, 26.98, 52.00, 58.93, 95.94, 183.85, 180.95, 186.21, 101.07, 178.49, 47.88, 92.91]
     return Elemets[i]

def list_ele_to_float(text):
    return float(text)

def mat_precent_ratio(list):
    tot = 0.0
    list_pre = []
    for i in range(len(list)):
       tot = tot+float(list[i])
    for i in range(len(list)):
       list_pre.append(str(round(float(list[i])/tot*100,2)))
    return list_pre
    
def mat_precent_minus(list):
    tot = 0.0
    list_pre = []
    for i in range(1, len(list)):
       tot = tot+float(list[i])
    list[0] = round((1.0 - tot/100)*100,2)
    for i in range(len(list)):
       list_pre.append(str(list[i]))
    return list_pre
    
def atom_precent(n,list):
    tot = 0.0
    list_pre = []
    atom_num = []
    for i in range(len(list)):
#        print(Atom_Element(num))
#        print(len)
        y =float(list[i])/float(Atom_Element(n,i))
        atom_num.append(str(y))
        tot = tot+ y
    for i in range(len(list)):
#        print (list.index(ele))
#        print (atom_num[list.index(ele)])
#        print(list[i], Atom_Element(i),atom_num[i])
        z=float(atom_num[i])/tot*100
#        print(tot)
        list_pre.append(str(round(z,2)))
    return list_pre

# print(1.0/Atom_Element(2))
def atom_num(list, Tot_num):
    num  = 0
    list_number = []
    for i in range(len(list)):
        z = Tot_num*float(list[i])/100
#        list_number.append(int(z+0.58))
        list_number.append(str(round(z,2)))
#    print(list_number)
    return list_number

#表格
tbs = doc.tables
n = 0
for tb in tbs:
    n = n+1
# 创建一行十列的表格(带边框)
#    table = document.add_table(rows=1, cols=14, style='Table Grid')
    table = document.add_table(rows=1, cols=12, style='Table Grid')
#    table.alignment = WD_TABLE_ALIGNMENT.CENTER #表格居中
    table.cell(0,0).width=Inches(5)
# 获取第一行的所有列数
    hdr_cells = table.rows[0].cells
# 给第一行的各个列添加内容
#    hdr_cells[0].text = ' No.'
#    hdr_cells[1].text = ' '
#    hdr_cells[2].text = ' Ni'
#    hdr_cells[3].text = ' Al'
#    hdr_cells[4].text = ' Cr'
#    hdr_cells[5].text = ' Co'
#    hdr_cells[6].text = ' Mo'
#    hdr_cells[7].text = ' W'
#    hdr_cells[8].text = ' Ta'
#    hdr_cells[9].text = ' Re'
#    hdr_cells[10].text = ' Ru'
#    hdr_cells[11].text = ' Hf'
#    hdr_cells[12].text = ' Ti'
#    hdr_cells[13].text = ' Nb'
    hdr_cells[0].text = ' No.'
    hdr_cells[1].text = ''
    hdr_cells[2].text = ' Ni'
    hdr_cells[3].text = ' Al'
    hdr_cells[4].text = ' Cr'
    hdr_cells[5].text = ' Co'
    hdr_cells[6].text = ' Mo'
    hdr_cells[7].text = ' W'
    hdr_cells[8].text = ' Ta'
    hdr_cells[9].text = ' Re'
    hdr_cells[10].text = ' Ru'
    if (n%3) == 2:
       hdr_cells[11].text = ' Ti'
    else:
       hdr_cells[11].text = ' Nb'

    #行
    rowline = 0
    for row in tb.rows[1:]:    
        #列    
        list_float=[]
        for cell in row.cells[1:]:
            try:
                x = float(cell.text)
                list_float.append(x)
            except TypeError:
                true
            except ValueError:
                if len(list_float):
                  x = 0.0
                  list_float.append(x)
                print(cell.text)

            #也可以用下面方法
            '''text = ''
            for p in cell.paragraphs:
                text += p.text
            print(text)'''

#        print(list_float)
# 给table表格添加新行，并给各列添加内容
#       for qty, id, desc in records:
        row_cells = table.add_row().cells
        rowline =rowline+1
        t = 0
        row_cells[t].text = ''
        row_cells[t+1].text = ' Wt%'
        for cell in row.cells[1:]:
            if float(list_float[t]) <= 0.0009:
                row_cells[t+2].text = ' '
            else:
                row_cells[t+2].text = str(list_float[t])
            t=t+1

#        mat=mat_precent_ratio(list_float)
        mat=mat_precent_minus(list_float)
#        print(mat)
# 给table表格添加新行，并给各列添加内容
#       for qty, id, desc in records:
        row_cells = table.add_row().cells
        rowline =rowline+1
        t = 0
        row_cells[t].text = ''
        row_cells[t+1].text = 'S-W%'
        table.cell(0,0).width=Inches(15) #设置单元格宽度
        for cell in row.cells[1:]:
            if float(mat[t]) <= 0.0009:
                row_cells[t+2].text = ' '
            else:
                row_cells[t+2].text = str(mat[t])
            t=t+1

        atom=atom_precent(n, mat)
#        print(atom)
        row_cells = table.add_row().cells
        rowline =rowline+1
        t = 0
        row_cells[t].text = ''
        row_cells[t+1].text = 'At%'
        for cell in row.cells[1:]:
            if float(atom[t]) <= 0.0009:
                row_cells[t+2].text = ' '
            else:
                row_cells[t+2].text = str(atom[t])
            t=t+1
#        mat=atom_precent(list_float)
#        print(mat)
        a_num=atom_num(atom, 392)
        row_cells = table.add_row().cells
        rowline =rowline+1
        t = 0
        row_cells[t].text = ''
        row_cells[t+1].text = 'Num'
        for cell in row.cells[1:]:
#            if int(a_num[t]) <= 0:
            if float(a_num[t]) <= 0.0009:
                row_cells[t+2].text = ' '
            else:
                row_cells[t+2].text = str(a_num[t])
            t=t+1

# 单元格合并 A[['列名1','列名2'....]].merge(B[['列名1','列名2'....]])

        table.cell(rowline-3,0).merge(table.cell(rowline,0)) #合并单元格(列)
#        table.cell(rowline,0).text='check' # 合并单元格后填入内容

        row_cells = table.add_row().cells
        rowline =rowline+1
#        print(rowline)
        table.cell(rowline,0).merge(table.cell(rowline,11)) #合并单元格(行)
#        row_cells[0].text = 'check'  #合并单元格后填入内容
        rate_Ni_Al = round(float(a_num[0])/float(a_num[1]),3)
        if (float(a_num[7])>= 0.00001 and float(a_num[8])>= 0.00001) :
          rate_Re_Ru = round(float(a_num[7])/float(a_num[8]),3)
          if float(rate_Re_Ru) <= 0.0001:
            rate_Ru_Re = ''
          else:
            rate_Ru_Re = round(float(1/rate_Re_Ru),2)
          rate_element = '      Ni : Al = '+ str(rate_Ni_Al) +' : 1 ,         Re : Ru = 1 : '+str(rate_Ru_Re)
        else:
          rate_element = '      Ni : Al = '+ str(rate_Ni_Al) +' : 1'
        row_cells[0].text = rate_element  #合并单元格后填入内容
        

#    document.paragraphs[0].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
    document.add_paragraph('')
# 添加分页
#    document.add_page_break()
#    document.paragraphs[0].runs[0].add_break()


#    row_cells[0].text = str(qty)
#    row_cells[1].text = id
#    row_cells[2].text = desc


# 保存world文档
#document.save('标准化数据表-2.docx')
document.save('标准化数据表-Wang_2021-04-06.docx')

if __name__ == "__main__":
    pass
